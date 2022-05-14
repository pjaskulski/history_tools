""" Skrypt do podziału pliku docx z tomem PSB na pliki
    z pojednycznymi biogramami
"""
import re
import docx
from docx import Document


def double_space(value:str) -> str:
    """ usuwa podwójne spacje z przekazanego tekstu """
    return ' '.join(value.strip().split())


def split(s: str) -> list:
    """ Dzieli tekst na części wg przecinka, pomijając przecinki wewnątrz
        nawiasów
        from: https://stackoverflow.com/questions/26808913/split-string-at-commas-
        except-when-in-bracket-environment/26809037
    """
    parts = []
    bracket_level = 0
    current = []
    # trick to remove special-case of trailing chars
    for c in (s + ","):
        if c == "," and bracket_level == 0:
            parts.append("".join(current))
            current = []
        else:
            if c == "(":
                bracket_level += 1
            elif c == ")":
                bracket_level -= 1
            current.append(c)
    return parts


nr = '27'
doc = docx.Document(f'psb_{nr}.docx')
lista = []

document = None
for par in doc.paragraphs:
    # if 'Schroeder' in par.text:
    #      print()
    new_doc = False
    if par.runs:
        if (par.runs[0].text == '\t' or par.runs[0].text.strip() == '') and len(par.runs) > 1:
            par_first = par.runs[1]
        else:
            par_first = par.runs[0]

        if (par_first.text.strip() != '' and par_first.bold is True
            or (par_first.bold is None and par_first.style
                and par_first.style.font.bold)):
            new_doc = True

    if not new_doc and not document:
        continue

    if new_doc:
        print(par_first.text.strip())
        if document:
            document.save(f'output/{nr}/{name.replace(" ", "_")}.docx')
        match = re.search(r'\d{4}', par.text)
        if len(par.text) < 100 and not match:
            document = None
            new_doc = False
            continue            # odsyłacze do innych haseł pomijamy

        document = Document()
        p = document.add_paragraph()
        for run in par.runs:
            output_run = p.add_run(run.text)
            output_run.bold = run.bold             # Run's bold data
            output_run.italic = run.italic         # Run's italic data
            output_run.underline = run.underline   # Run's underline data
            output_run.font.color.rgb = run.font.color.rgb # Run's color data
            if run.style:
                output_run.style.font.bold = run.style.font.bold
                output_run.style.font.italic = run.style.font.italic


        #tmp = par.text.strip().split(",")
        tmp = split(par.text.strip())

        name = tmp[0].strip()
        if "(" in name:
            pattern = r'[\(\[].*?[\)\]]'
            name = re.sub(pattern, '', name)
            name = double_space(name)
            name = name.strip()

        tmp = name.split(" ")
        if len(tmp) > 3:
            tmp = tmp[:3]
            name = ' '.join(tmp)
            if name.endswith('_ze'):
                name = name[:-3]
            elif name.endswith('_z'):
                name = name[:-2]

        if name not in lista:
            lista.append(name)
        else:
            licznik = 0
            while True:
                licznik += 1
                new_name = f'{name}_{str(licznik)}'
                if new_name not in lista:
                    lista.append(new_name)
                    name = new_name
                    break

    else:
        p = document.add_paragraph()
        for run in par.runs:
            output_run = p.add_run(run.text)
            output_run.bold = run.bold           # Run's bold data
            output_run.italic = run.italic       # Run's italic data
            output_run.underline = run.underline # Run's underline data
            output_run.font.name = run.font.name
            output_run.font.color.rgb = run.font.color.rgb # Run's color data
            output_run.font.italic = run.font.italic # Run's font
            if run.style:
                output_run.style.font.italic = run.style.font.italic


        # Paragraph's alignment data
        p.alignment = par.alignment
        p.paragraph_format.alignment = par.paragraph_format.alignment

if document:
    document.save(f'output/{nr}/{name.replace(" ", "_")}.docx')
